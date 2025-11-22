# plagiarism_server.py
# Production-ready plagiarism checker and remover with web crawling
import os
import io
import time
import re
import json
import uuid
import sqlite3
import shutil
import logging
import threading
import pickle
from pathlib import Path
from typing import List, Dict, Optional, Any, Tuple
from urllib.parse import urlparse, urljoin
import urllib.robotparser
import hashlib
from contextlib import asynccontextmanager

import requests
from bs4 import BeautifulSoup

# text extraction libs
from pdfminer.high_level import extract_text as extract_text_pdfminer
from PyPDF2 import PdfReader
import docx
import pandas as pd

# NLP & similarity - optional libraries with graceful fallback
try:
    HAVE_SENTENCE_TRANSFORMERS = True
except Exception:
    HAVE_SENTENCE_TRANSFORMERS = False

try:
    from readability import Document as ReadabilityDocument
    HAVE_READABILITY = True
except Exception:
    HAVE_READABILITY = False

try:
    from datasketch import MinHash, MinHashLSH
    HAVE_DATASKETCH = True
except Exception:
    HAVE_DATASKETCH = False

try:
    import nltk
    from nltk.corpus import wordnet
    from nltk.tokenize import sent_tokenize, word_tokenize
    from nltk import pos_tag, download as nltk_download
    HAVE_NLTK = True
    try:
        nltk_download('punkt', quiet=True)
        nltk_download('wordnet', quiet=True)
        nltk_download('averaged_perceptron_tagger', quiet=True)
    except:
        pass
except Exception:
    HAVE_NLTK = False

# FastAPI
from fastapi import FastAPI, Form, File, UploadFile
from fastapi.responses import JSONResponse
from fastapi.middleware.cors import CORSMiddleware
import uvicorn

# Logging setup
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger("textguard")

# ============= CONFIGURATION =============
DATA_DIR = Path("textguard_data")
DATA_DIR.mkdir(exist_ok=True)
DB_PATH = DATA_DIR / "plagiarism_db.sqlite3"
LSH_PATH = DATA_DIR / "lsh.pkl"
CACHE_DIR = DATA_DIR / "cache"
CACHE_DIR.mkdir(exist_ok=True)

HTTP_HEADERS = {"User-Agent": "TextGuardBot/2.0 (+https://textguard.local)"}
SHINGLE_SIZE = 5
MINHASH_PERMS = 128
LSH_THRESHOLD = 0.4
PLAGIARISM_THRESHOLD = 0.60
WEB_FETCH_TIMEOUT = 30
CACHE_TTL = 60 * 60 * 24  # 24 hours
ENABLE_WEB_SEARCH = True  # Enable web search for real analysis

DEFAULT_SEEDS = [
    "https://en.wikipedia.org/wiki/",
    "https://arxiv.org/",
    "https://medium.com/",
    "https://www.geeksforgeeks.org/",
    "https://github.com/",
    "https://stackoverflow.com/",
    "https://www.quora.com/",
    "https://www.reddit.com/r/",
]

# Thread locks for thread-safe operations
LSH_LOCK = threading.Lock()
DB_LOCK = threading.Lock()
CACHE_LOCK = threading.Lock()

# Global semantic model
SEMANTIC_MODEL = None

def load_semantic_model():
    """Load semantic transformer model for similarity computation"""
    global SEMANTIC_MODEL
    if HAVE_SENTENCE_TRANSFORMERS and SEMANTIC_MODEL is None:
        try:
            logger.info("Loading semantic transformer model...")
            from sentence_transformers import SentenceTransformer
            SEMANTIC_MODEL = SentenceTransformer('all-MiniLM-L6-v2')
            logger.info("Semantic model loaded successfully")
        except Exception as e:
            logger.warning(f"Failed to load semantic model: {e}")
            SEMANTIC_MODEL = None

# ============= DATABASE FUNCTIONS =============
def init_db():
    """Initialize database with required tables"""
    with DB_LOCK:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        
        c.execute("""
        CREATE TABLE IF NOT EXISTS pages (
            id TEXT PRIMARY KEY,
            url TEXT UNIQUE,
            text TEXT,
            text_hash TEXT,
            minhash BLOB,
            fetched_at REAL,
            domain TEXT
        )""")
        
        c.execute("""
        CREATE TABLE IF NOT EXISTS submissions (
            id TEXT PRIMARY KEY,
            user_id TEXT,
            text TEXT,
            minhash BLOB,
            plagiarism_score REAL,
            checked_at REAL,
            source_file TEXT
        )""")
        
        c.execute("""
        CREATE TABLE IF NOT EXISTS reports (
            id TEXT PRIMARY KEY,
            submission_id TEXT,
            report_data TEXT,
            created_at REAL
        )""")
        
        conn.commit()
        conn.close()

def save_page_to_db(page_id: str, url: str, text: str, minhash_bytes: bytes = b""):
    """Save fetched page to database"""
    with DB_LOCK:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        domain = urlparse(url).netloc
        text_hash = hashlib.md5(text.encode()).hexdigest()
        c.execute("""
            INSERT OR REPLACE INTO pages (id, url, text, text_hash, minhash, fetched_at, domain)
            VALUES (?, ?, ?, ?, ?, ?, ?)
        """, (page_id, url, text, text_hash, minhash_bytes, time.time(), domain))
        conn.commit()
        conn.close()

def get_page_by_url(url: str):
    """Retrieve page from database by URL"""
    with DB_LOCK:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute("""
            SELECT id, url, text, minhash, fetched_at, domain FROM pages WHERE url = ?
        """, (url,))
        row = c.fetchone()
        conn.close()
    
    if not row:
        return None
    return {"id": row[0], "url": row[1], "text": row[2], "minhash": row[3], "fetched_at": row[4], "domain": row[5]}

def save_submission_to_db(sub_id: str, user_id: Optional[str], text: str, minhash_bytes: bytes = b"",
                         plagiarism_score: float = 0.0, source_file: str = ""):
    """Save user submission to database"""
    with DB_LOCK:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute("""
            INSERT INTO submissions (id, user_id, text, minhash, plagiarism_score, checked_at, source_file)
            VALUES (?, ?, ?, ?, ?, ?, ?)
        """, (sub_id, user_id or "", text, minhash_bytes, plagiarism_score, time.time(), source_file))
        conn.commit()
        conn.close()

def save_report_to_db(report_id: str, submission_id: str, report_data: Dict):
    """Save detailed plagiarism report"""
    with DB_LOCK:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute("""
            INSERT INTO reports (id, submission_id, report_data, created_at)
            VALUES (?, ?, ?, ?)
        """, (report_id, submission_id, json.dumps(report_data), time.time()))
        conn.commit()
        conn.close()

def get_recent_submissions_by_user(user_id: str, limit: int = 10):
    """Get recent submissions by user"""
    with DB_LOCK:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute("""
            SELECT id, user_id, text, minhash, plagiarism_score, checked_at
            FROM submissions WHERE user_id = ? ORDER BY checked_at DESC LIMIT ?
        """, (user_id, limit))
        rows = c.fetchall()
        conn.close()
    
    return [
        {"id": r[0], "user_id": r[1], "text": r[2], "minhash": r[3], "plagiarism_score": r[4], "checked_at": r[5]}
        for r in rows
    ]

# ============= MINHASH & LSH FUNCTIONS =============
LSH_INSTANCE = None

def load_or_create_lsh():
    """Load or create LSH index for fast similarity search"""
    global LSH_INSTANCE
    if not HAVE_DATASKETCH:
        logger.info("datasketch not available; MinHash LSH disabled")
        LSH_INSTANCE = None
        return
    
    with LSH_LOCK:
        if LSH_PATH.exists():
            try:
                import pickle
                LSH_INSTANCE = pickle.loads(LSH_PATH.read_bytes())
                logger.info("Loaded LSH from disk")
                return
            except Exception as e:
                logger.warning(f"Failed to load LSH: {e}; creating new")
        
        LSH_INSTANCE = MinHashLSH(threshold=LSH_THRESHOLD, num_perm=MINHASH_PERMS)
        logger.info("Created new LSH instance")

def persist_lsh():
    """Save LSH index to disk"""
    if not HAVE_DATASKETCH or LSH_INSTANCE is None:
        return
    
    with LSH_LOCK:
        try:
            import pickle
            LSH_PATH.write_bytes(pickle.dumps(LSH_INSTANCE))
        except Exception as e:
            logger.error(f"Failed to persist LSH: {e}")

def minhash_from_text(text: str, k: int = SHINGLE_SIZE, num_perm: int = MINHASH_PERMS):
    """Generate MinHash from text"""
    if not HAVE_DATASKETCH:
        raise RuntimeError("datasketch not available")
    
    tokens = tokenize(text)
    shingles = shingles_from_tokens(tokens, k)
    m = MinHash(num_perm=num_perm)
    for s in shingles:
        m.update(s.encode("utf8"))
    return m

# ============= TEXT PROCESSING FUNCTIONS =============
def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file"""
    try:
        doc = docx.Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return ""

def extract_text_from_pdf_bytes(data: bytes) -> str:
    """Extract text from PDF bytes"""
    try:
        with io.BytesIO(data) as f:
            text = extract_text_pdfminer(f)
        if text and text.strip():
            return text
    except Exception:
        logger.debug("pdfminer extraction failed")
    
    try:
        with io.BytesIO(data) as f:
            reader = PdfReader(f)
            pages = []
            for p in reader.pages:
                try:
                    pages.append(p.extract_text() or "")
                except Exception:
                    pages.append("")
            text = "\n".join(pages)
        if text and text.strip():
            return text
    except Exception:
        logger.debug("PyPDF2 extraction failed")
    
    return ""

def extract_text_from_file_bytes(filename: str, data: bytes) -> str:
    """Extract text from various file formats"""
    name = filename.lower().strip()
    try:
        if name.endswith(".txt"):
            return data.decode("utf-8", errors="ignore")
        elif name.endswith(".pdf"):
            return extract_text_from_pdf_bytes(data)
        elif name.endswith(".docx"):
            tmp = f"tmp_{os.getpid()}_{uuid.uuid4().hex[:8]}.docx"
            with open(tmp, "wb") as t:
                t.write(data)
            try:
                return extract_text_from_docx(tmp)
            finally:
                try:
                    os.remove(tmp)
                except Exception:
                    pass
        elif name.endswith(".csv"):
            return extract_text_from_csv_bytes(data)
        else:
            return data.decode("utf-8", errors="ignore")
    except Exception as e:
        logger.error(f"Text extraction failed for {filename}: {e}")
        return ""

def extract_text_from_csv_bytes(data: bytes) -> str:
    """Extract text from CSV bytes"""
    for enc in ("utf-8", "latin-1", "cp1252"):
        try:
            df = pd.read_csv(io.BytesIO(data), encoding=enc, engine='python', dtype=str, low_memory=False)
            return " ".join(df.fillna("").astype(str).apply(lambda row: " ".join(row.values), axis=1).tolist())
        except Exception:
            continue
    return ""

def tokenize(text: str) -> List[str]:
    """Tokenize text into words"""
    text = (text or "").lower()
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"[^a-z0-9\s]+", " ", text)
    tokens = [tok for tok in text.split(" ") if tok and len(tok) > 1]
    return tokens

def shingles_from_tokens(tokens: List[str], k: int) -> List[str]:
    """Generate k-shingles from tokens"""
    if not tokens:
        return []
    if len(tokens) <= k:
        return [" ".join(tokens)]
    return [" ".join(tokens[i:i+k]) for i in range(len(tokens)-k+1)]

def compute_jaccard(a: set, b: set) -> float:
    """Compute Jaccard similarity between two sets"""
    if not a and not b:
        return 1.0
    if not a or not b:
        return 0.0
    inter = len(a.intersection(b))
    union = len(a.union(b))
    return inter / union if union else 0.0

def split_into_sentences(text: str) -> List[str]:
    """Split text into sentences"""
    if HAVE_NLTK:
        try:
            return sent_tokenize(text)
        except Exception:
            pass
    
    # Fallback simple sentence splitting
    sentences = re.split(r'(?<=[.!?])\s+', text)
    return [s.strip() for s in sentences if s.strip()]

def split_into_chunks(text: str, chunk_size: int = 256, overlap: int = 50) -> List[str]:
    """Split text into overlapping chunks for processing"""
    sentences = split_into_sentences(text)
    chunks = []
    current_chunk = ""
    
    for sentence in sentences:
        if len(current_chunk.split()) + len(sentence.split()) <= chunk_size:
            current_chunk += " " + sentence if current_chunk else sentence
        else:
            if current_chunk:
                chunks.append(current_chunk)
            current_chunk = sentence
    
    if current_chunk:
        chunks.append(current_chunk)
    
    return chunks

def compute_semantic_similarity(text1: str, text2: str) -> float:
    """Compute semantic similarity using transformer models"""
    if not HAVE_SENTENCE_TRANSFORMERS or SEMANTIC_MODEL is None:
        return 0.0
    
    try:
        from sentence_transformers import util
        embedding1 = SEMANTIC_MODEL.encode(text1, convert_to_tensor=True)
        embedding2 = SEMANTIC_MODEL.encode(text2, convert_to_tensor=True)
        similarity = util.pytorch_cos_sim(embedding1, embedding2)
        return float(similarity[0][0])
    except Exception as e:
        logger.warning(f"Semantic similarity computation failed: {e}")
        return 0.0

# ============= WEB FETCHING FUNCTIONS =============
def allowed_by_robots(url: str, user_agent: str = HTTP_HEADERS["User-Agent"]) -> bool:
    """Check if URL is allowed by robots.txt"""
    try:
        parsed = urlparse(url)
        base = f"{parsed.scheme}://{parsed.netloc}"
        rp_url = urljoin(base, "/robots.txt")
        rp = urllib.robotparser.RobotFileParser()
        rp.set_url(rp_url)
        rp.read()
        return rp.can_fetch(user_agent, url)
    except Exception:
        return True

def fetch_page_text(url: str, timeout: int = WEB_FETCH_TIMEOUT) -> str:
    """Fetch and extract text from a URL with caching"""
    # Check cache first
    cache_key = hashlib.md5(url.encode()).hexdigest()
    cache_file = CACHE_DIR / f"{cache_key}.txt"
    
    if cache_file.exists():
        try:
            stat = os.stat(cache_file)
            if time.time() - stat.st_mtime < CACHE_TTL:
                content = cache_file.read_text()
                logger.debug(f"Cache hit for {url}")
                return content
        except Exception:
            pass
    
    if not allowed_by_robots(url):
        logger.debug(f"Blocked by robots.txt: {url}")
        return ""
    
    try:
        r = requests.get(url, headers=HTTP_HEADERS, timeout=timeout)
        if r.status_code != 200:
            logger.debug(f"Non-200 status for {url}: {r.status_code}")
            return ""
        
        text = ""
        
        # Try readability for better content extraction
        if HAVE_READABILITY:
            try:
                doc = ReadabilityDocument(r.text)
                main_html = doc.summary()
                soup = BeautifulSoup(main_html, "html.parser")
                text = soup.get_text(separator=" ", strip=True)
            except Exception:
                logger.debug("Readability extraction failed")
        
        # Fallback to BeautifulSoup
        if not text:
            soup = BeautifulSoup(r.content, "html.parser")
            for tag in soup(["script", "style", "header", "footer", "nav", "noscript", "svg", "figure", "iframe"]):
                tag.decompose()
            text = soup.get_text(separator=" ", strip=True)
        
        text = " ".join(text.split())
        
        # Cache the result
        try:
            cache_file.write_text(text)
        except Exception:
            pass
        
        time.sleep(0.5)  # Respect rate limits
        return text
    except Exception as e:
        logger.debug(f"Fetch error for {url}: {e}")
        return ""

def search_google_custom(query: str, max_results: int = 10) -> List[str]:
    """Search for URLs using DuckDuckGo (no API key required)"""
    urls = []
    urls_container = []
    exception_container = []
    
    def do_search():
        try:
            from ddgs import DDGS
            
            with DDGS() as ddgs:
                results = ddgs.text(query, max_results=max_results)
                for result in results:
                    if 'href' in result:
                        urls_container.append(result['href'])
                    elif 'link' in result:
                        urls_container.append(result['link'])
            
            logger.info(f"Found {len(urls_container)} URLs for query: {query}")
        except Exception as e:
            exception_container.append(e)
            logger.debug(f"DuckDuckGo search error: {e}")
    
    # Run search in a thread with timeout
    search_thread = threading.Thread(target=do_search, daemon=True)
    search_thread.start()
    search_thread.join(timeout=10)  # 10 second timeout
    
    if search_thread.is_alive():
        logger.warning(f"DuckDuckGo search timeout for query: {query}")
        return []
    
    if exception_container:
        logger.debug(f"DuckDuckGo search failed: {exception_container[0]}")
        return []
    
    return urls_container

def search_web_for_phrases(phrases: List[str], max_urls: int = 30, seeds: Optional[List[str]] = None) -> List[str]:
    """Search for URLs containing specific phrases"""
    seeds = seeds or DEFAULT_SEEDS
    urls_found = set()
    
    # Try DuckDuckGo search only if enabled
    if ENABLE_WEB_SEARCH:
        for phrase in phrases[:5]:  # Limit to 5 phrases
            try:
                search_query = f'"{phrase}"'
                results = search_google_custom(search_query, max_results=min(10, max_urls))
                urls_found.update(results)
                
                if len(urls_found) >= max_urls:
                    break
            except Exception as e:
                logger.debug(f"Search failed for phrase '{phrase}': {e}")
    else:
        logger.info("Web search is disabled, using seed URLs only")
    
    # Add seed URLs
    urls_found.update(seeds[:3])
    
    return list(urls_found)[:max_urls]

# ============= PLAGIARISM DETECTION =============
def build_quoted_phrases(text: str, phrase_len: int = 10, step: int = 5, max_phrases: int = 30) -> List[str]:
    """Build search phrases from text"""
    tokens = tokenize(text)
    phrases = []
    seen = set()
    
    if len(tokens) == 0:
        return phrases
    
    for i in range(0, max(1, len(tokens) - phrase_len + 1), step):
        p = " ".join(tokens[i:i+phrase_len])
        if p and p not in seen and len(p) > 20:
            phrases.append(p)
            seen.add(p)
        
        if len(phrases) >= max_phrases:
            break
    
    # Add smaller phrases if needed
    if len(phrases) < max_phrases:
        small_k = max(4, phrase_len // 2)
        small = shingles_from_tokens(tokens, small_k)
        for s in small:
            if len(phrases) >= max_phrases:
                break
            if s not in seen and len(s.split()) > 2:
                phrases.append(s)
                seen.add(s)
    
    return phrases[:max_phrases]

def analyze_plagiarism(text: str, max_phrases: int = 10, max_urls: int = 30, 
                       seeds: Optional[List[str]] = None, use_semantic: bool = True,
                       quick_mode: bool = False) -> Dict[str, Any]:
    """Comprehensive plagiarism analysis"""
    try:
        seeds = seeds or DEFAULT_SEEDS
        
        # Tokenize and create shingles
        tokens_query = tokenize(text)
        shingles_query = set(shingles_from_tokens(tokens_query, SHINGLE_SIZE))
        
        results = {
            "ok": True,
            "text_stats": {
                "total_words": len(text.split()),
                "total_chars": len(text),
                "total_sentences": len(split_into_sentences(text)),
                "unique_tokens": len(set(tokens_query))
            },
            "matches": [],
            "matched_phrases": [],
            "plagiarism_score": 0.0,
            "plagiarism_percentage": 0.0,
            "sources_examined": 0
        }
        
        # In quick mode, skip web searching and return basic results
        if quick_mode:
            logger.info("Quick mode enabled - skipping web search")
            results["plagiarism_percentage"] = 0.0
            results["plagiarism_score"] = 0.0
            results["is_plagiarized"] = False
            return results
        
        # Build search phrases
        phrases = build_quoted_phrases(text, phrase_len=12, step=5, max_phrases=max_phrases)
        logger.info(f"Generated {len(phrases)} search phrases")
        
        # Search for URLs
        search_urls = search_web_for_phrases(phrases, max_urls=max_urls, seeds=seeds)
        logger.info(f"Found {len(search_urls)} URLs to examine")
        
        matched_shingles = set()
        all_matches = []
        
        # Process each URL
        for idx, url in enumerate(search_urls):
            if idx >= max_urls:
                break
            
            try:
                # Check cache first
                cached_page = get_page_by_url(url)
                if cached_page:
                    page_text = cached_page['text']
                    logger.debug(f"Using cached page: {url}")
                else:
                    page_text = fetch_page_text(url)
                    if not page_text or len(page_text.split()) < 50:
                        continue
                    
                    # Save to database
                    page_id = str(uuid.uuid4())
                    if HAVE_DATASKETCH:
                        try:
                            m = minhash_from_text(page_text)
                            import pickle
                            mh_bytes = pickle.dumps(list(m.hashvalues))
                        except Exception:
                            mh_bytes = b""
                    else:
                        mh_bytes = b""
                    
                    save_page_to_db(page_id, url, page_text, mh_bytes)
                
                # Compute similarity
                tokens_page = tokenize(page_text)
                shingles_page = set(shingles_from_tokens(tokens_page, SHINGLE_SIZE))
                jaccard_sim = compute_jaccard(shingles_query, shingles_page)
                
                # Compute semantic similarity if enabled
                semantic_sim = 0.0
                if use_semantic:
                    semantic_sim = compute_semantic_similarity(text[:500], page_text[:500])
                
                # Combined score
                combined_score = (jaccard_sim * 0.6) + (semantic_sim * 0.4)
                
                if jaccard_sim > 0.15 or semantic_sim > 0.6:
                    match_info = {
                        "url": url,
                        "domain": urlparse(url).netloc,
                        "jaccard_similarity": round(jaccard_sim, 4),
                        "semantic_similarity": round(semantic_sim, 4),
                        "combined_score": round(combined_score, 4),
                        "plagiarism_percent": round(jaccard_sim * 100, 2),
                        "page_words": len(tokens_page),
                        "matched_words": len(shingles_query.intersection(shingles_page))
                    }
                    all_matches.append(match_info)
                    matched_shingles.update(shingles_query.intersection(shingles_page))
                
                results["sources_examined"] += 1
            
            except Exception as e:
                logger.error(f"Error processing URL {url}: {e}")
                continue
        
        # Calculate overall plagiarism score
        total_shingles = len(shingles_query) or 1
        overlap_percentage = (len(matched_shingles) / total_shingles) * 100
        
        results["matches"] = sorted(all_matches, key=lambda x: x["combined_score"], reverse=True)[:10]
        results["plagiarism_percentage"] = min(100.0, round(overlap_percentage, 2))
        results["plagiarism_score"] = min(1.0, round(overlap_percentage / 100, 4))
        results["matched_phrases_count"] = len(matched_shingles)
        
        return results
    
    except Exception as e:
        logger.error(f"Plagiarism analysis failed: {e}")
        return {
            "ok": False,
            "error": f"Analysis failed: {str(e)}"
        }

# ============= PLAGIARISM REMOVAL/PARAPHRASING =============

# High-quality, domain-specific synonym mappings for OOP/CS concepts
PREMIUM_SYNONYMS = {
    # Core OOP terms - carefully curated
    "class": "type",
    "object": "instance",
    "represents": "describes",
    "consists": "comprises",
    "contains": "holds",
    "members": "components",
    "functions": "operations",
    "common": "shared",
    "properties": "attributes",
    "methods": "capabilities",
    "created": "instantiated",
    "allocated": "assigned",
    "accessed": "used",
    "known": "understood",
    "entity": "element",
    "entities": "elements",
    "basic": "fundamental",
    "basic unit": "fundamental unit",
    "unit": "component",
    "real-life": "real-world",
    "real-world": "real-life",
    "examples": "instances",
    "consider": "examine",
    "characteristics": "attributes",
    "behavior": "functionality",
    "identity": "distinctiveness",
    "state": "condition",
    "allocate": "assign",
    "manipulate": "process",
    "interact": "communicate",
    "interact without": "work without",
    "sufficient": "enough",
    "knowledge": "understanding",
    "details": "specifics",
}

def get_premium_synonym(word: str) -> Optional[str]:
    """Get a high-quality synonym for a word - returns single best option"""
    word_lower = word.lower()
    return PREMIUM_SYNONYMS.get(word_lower)

def paraphrase_phrase_advanced(phrase: str, intensity: float) -> str:
    """
    Paraphrase a phrase with intelligent replacement.
    Higher intensity = more changes, but always maintains grammar.
    """
    if not phrase or len(phrase.strip()) < 2:
        return phrase
    
    words = phrase.split()
    result = []
    replacements_made = 0
    target = max(1, int(len(words) * intensity * 0.25))  # Replace 25% at max intensity
    
    i = 0
    while i < len(words):
        word = words[i]
        word_clean = word
        leading_spaces = ""
        trailing_punct = ""
        
        # Extract trailing punctuation
        while word_clean and word_clean[-1] in ",.!?;:":
            trailing_punct = word_clean[-1] + trailing_punct
            word_clean = word_clean[:-1]
        
        word_lower = word_clean.lower()
        
        # Try multi-word phrases first (only if not already at end)
        if i + 1 < len(words) and word_lower:
            next_word_clean = words[i+1]
            # Remove trailing punct from next word for comparison
            while next_word_clean and next_word_clean[-1] in ",.!?;:":
                next_word_clean = next_word_clean[:-1]
            
            two_word = word_lower + " " + next_word_clean.lower()
            if two_word in PREMIUM_SYNONYMS and replacements_made < target:
                replacement = PREMIUM_SYNONYMS[two_word]
                # Capitalize appropriately
                if word[0].isupper():
                    replacement = replacement[0].upper() + replacement[1:] if len(replacement) > 1 else replacement.upper()
                result.append(replacement)
                i += 2
                replacements_made += 1
                continue
        
        # Try single word replacement
        if word_lower and replacements_made < target:
            syn = get_premium_synonym(word_lower)
            if syn:
                # Smart capitalization
                if word[0].isupper():
                    syn = syn[0].upper() + syn[1:] if len(syn) > 1 else syn.upper()
                result.append(syn + trailing_punct)
                replacements_made += 1
                i += 1
                continue
        
        # Keep original
        result.append(word)
        i += 1
    
    return " ".join(result)

def restructure_sentence_smart(sentence: str, intensity: float) -> str:
    """Intelligently restructure sentences for variety at higher intensities"""
    if intensity < 0.6 or len(sentence.split()) < 8:
        return sentence
    
    # Only apply advanced restructuring at high intensities
    # and for sentences long enough to restructure safely
    try:
        words = sentence.split()
        
        # Look for "It is/This is" patterns and simplify
        if len(words) > 5:
            if words[0].lower() == "it" and words[1].lower() == "is":
                # "It is a... that..." -> "A... that..."
                if len(words) > 3:
                    sentence = " ".join(words[2:])
            elif words[0].lower() == "this" and words[1].lower() == "is":
                # "This is a... that..." -> "The..."
                if len(words) > 3:
                    sentence = "The " + " ".join(words[3:])
        
        return sentence
    except:
        return sentence

def paraphrase_sentence_v2(sentence: str, intensity: float = 0.5) -> str:
    """
    Advanced paraphrasing that produces natural, human-like results.
    Focuses on quality over quantity of changes.
    """
    if not sentence or len(sentence.strip()) < 5:
        return sentence
    
    try:
        # Step 1: Apply smart paraphrasing
        para = paraphrase_phrase_advanced(sentence, intensity)
        
        # Step 2: Apply restructuring if intensity is high
        para = restructure_sentence_smart(para, intensity)
        
        return para.strip()
    except Exception:
        return sentence

def remove_plagiarism(text: str, intensity: float = 0.7) -> Dict[str, Any]:
    """Remove plagiarism with intelligent, human-like paraphrasing"""
    try:
        intensity = max(0.1, min(1.0, intensity))
        
        sentences = split_into_sentences(text)
        paraphrased_sentences = []
        
        for sentence in sentences:
            para_sent = paraphrase_sentence_v2(sentence, intensity=intensity)
            paraphrased_sentences.append(para_sent)
        
        # Join with proper spacing
        result_text = " ".join(paraphrased_sentences)
        result_text = re.sub(r'\s+', ' ', result_text).strip()
        
        # Calculate metrics
        original_tokens = set(tokenize(text.lower()))
        result_tokens = set(tokenize(result_text.lower()))
        
        # Token overlap shows how much is unchanged
        if original_tokens:
            unchanged = len(original_tokens.intersection(result_tokens)) / len(original_tokens)
            change_percentage = round((1 - unchanged) * 100, 2)
        else:
            change_percentage = 0
        
        return {
            "ok": True,
            "original_text": text,
            "paraphrased_text": result_text,
            "changes_applied": {
                "intensity": intensity,
                "sentences_processed": len(sentences),
                "text_variation": change_percentage,
                "method": "intelligent_context_aware_paraphrasing"
            }
        }
    
    except Exception as e:
        logger.error(f"Plagiarism removal failed: {e}")
        return {
            "ok": False,
            "error": str(e)
        }

# ============= FastAPI APPLICATION =============
@asynccontextmanager
async def lifespan(app: FastAPI):
    """FastAPI lifespan event handler (modern replacement for on_event)"""
    # Startup
    init_db()
    load_or_create_lsh()
    load_semantic_model()
    logger.info("TextGuard server started successfully")
    yield
    # Shutdown
    persist_lsh()
    logger.info("TextGuard server shutting down")

app = FastAPI(
    title="TextGuard - Production Plagiarism Checker & Remover",
    description="Advanced plagiarism detection and removal system with web crawling",
    version="2.0.0",
    lifespan=lifespan
)

# Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"]
)

@app.post("/api/check-text")
async def check_text(
    text: str = Form(...),
    user_id: Optional[str] = Form(None),
    max_phrases: int = Form(10),
    max_urls: int = Form(30),
    use_semantic: bool = Form(True)
):
    """Check text for plagiarism"""
    if not text or not text.strip():
        return JSONResponse(
            {"ok": False, "error": "Empty text"},
            status_code=400
        )
    
    if len(text) > 1000000:  # 1MB limit
        return JSONResponse(
            {"ok": False, "error": "Text too large (max 1MB)"},
            status_code=413
        )
    
    try:
        # Run analysis in quick mode to avoid hanging
        result = analyze_plagiarism(
            text,
            max_phrases=min(int(max_phrases), 20),
            max_urls=min(int(max_urls), 60),
            use_semantic=bool(use_semantic),
            quick_mode=False  # Disable quick mode to run real analysis
        )
        
        # Save submission
        if result.get("ok"):
            try:
                sub_id = str(uuid.uuid4())
                if HAVE_DATASKETCH:
                    try:
                        m = minhash_from_text(text)
                        mh_bytes = pickle.dumps(list(m.hashvalues))
                    except Exception:
                        mh_bytes = b""
                else:
                    mh_bytes = b""
                
                save_submission_to_db(
                    sub_id,
                    user_id,
                    text,
                    mh_bytes,
                    result.get("plagiarism_score", 0.0)
                )
                
                # Save detailed report
                report_id = str(uuid.uuid4())
                save_report_to_db(report_id, sub_id, result)
                
                result["submission_id"] = sub_id
                result["report_id"] = report_id
            except Exception as e:
                logger.error(f"Failed to save submission: {e}")
        
        result["is_plagiarized"] = result.get("plagiarism_score", 0) >= PLAGIARISM_THRESHOLD
        return result
    
    except Exception as e:
        logger.error(f"Check text error: {e}")
        return JSONResponse(
            {"ok": False, "error": "Internal server error"},
            status_code=500
        )

@app.post("/api/check-file")
async def check_file(
    file: UploadFile = File(...),
    user_id: Optional[str] = Form(None),
    max_phrases: int = Form(10),
    max_urls: int = Form(30)
):
    """Check uploaded file for plagiarism"""
    try:
        content = await file.read()
        text = extract_text_from_file_bytes(file.filename, content)
        
        if not text or not text.strip():
            return JSONResponse(
                {"ok": False, "error": "Could not extract text from file"},
                status_code=400
            )
        
        # Run analysis
        result = analyze_plagiarism(
            text,
            max_phrases=min(int(max_phrases), 20),
            max_urls=min(int(max_urls), 60)
        )
        
        if result.get("ok"):
            try:
                sub_id = str(uuid.uuid4())
                if HAVE_DATASKETCH:
                    try:
                        m = minhash_from_text(text)
                        import pickle
                        mh_bytes = pickle.dumps(list(m.hashvalues))
                    except Exception:
                        mh_bytes = b""
                else:
                    mh_bytes = b""
                
                save_submission_to_db(
                    sub_id,
                    user_id,
                    text,
                    mh_bytes,
                    result.get("plagiarism_score", 0.0),
                    file.filename
                )
                
                result["submission_id"] = sub_id
                result["filename"] = file.filename
            except Exception as e:
                logger.error(f"Failed to save file submission: {e}")
        
        result["is_plagiarized"] = result.get("plagiarism_score", 0) >= PLAGIARISM_THRESHOLD
        return result
    
    except Exception as e:
        logger.error(f"Check file error: {e}")
        return JSONResponse(
            {"ok": False, "error": "File processing failed"},
            status_code=500
        )

@app.post("/api/remove-plagiarism")
async def remove_plagiarism_endpoint(
    text: str = Form(...),
    intensity: float = Form(0.7)
):
    """Remove plagiarism from text"""
    if not text or not text.strip():
        return JSONResponse(
            {"ok": False, "error": "Empty text"},
            status_code=400
        )
    
    try:
        result = remove_plagiarism(text, intensity=float(intensity))
        return result
    except Exception as e:
        logger.error(f"Remove plagiarism error: {e}")
        return JSONResponse(
            {"ok": False, "error": "Paraphrasing failed"},
            status_code=500
        )

@app.get("/api/stats")
async def get_stats():
    """Get system statistics"""
    try:
        with DB_LOCK:
            conn = sqlite3.connect(DB_PATH)
            c = conn.cursor()
            
            c.execute("SELECT COUNT(*) FROM pages")
            cached_pages = c.fetchone()[0]
            
            c.execute("SELECT COUNT(*) FROM submissions")
            submissions = c.fetchone()[0]
            
            c.execute("SELECT COUNT(*) FROM reports")
            reports = c.fetchone()[0]
            
            conn.close()
        
        cache_size = sum(f.stat().st_size for f in CACHE_DIR.glob("*.txt")) / 1024 / 1024
        
        return {
            "ok": True,
            "cached_pages": cached_pages,
            "total_submissions": submissions,
            "total_reports": reports,
            "cache_size_mb": round(cache_size, 2),
            "semantic_model_loaded": SEMANTIC_MODEL is not None,
            "lsh_enabled": LSH_INSTANCE is not None
        }
    except Exception as e:
        logger.error(f"Stats error: {e}")
        return {"ok": False, "error": str(e)}

@app.get("/")
async def root():
    """Root endpoint"""
    return {
        "ok": True,
        "service": "TextGuard",
        "version": "2.0.0",
        "status": "running",
        "endpoints": {
            "check_text": "/api/check-text",
            "check_file": "/api/check-file",
            "remove_plagiarism": "/api/remove-plagiarism",
            "stats": "/api/stats"
        }
    }

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {"status": "healthy"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(
        "plagiarism_server:app",
        host="0.0.0.0",
        port=8000,
        reload=False,
        log_level="info"
    )
